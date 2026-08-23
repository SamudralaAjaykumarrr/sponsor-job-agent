from pathlib import Path

from docx import Document
from docx.shared import Pt

from app.resume.generator import ResumeContent

_MAX_LEVEL = 6
_BASE_FONT_SIZE = 10.5


def _clamp(level: int) -> int:
    return max(0, min(level, _MAX_LEVEL))


def write_docx(resume: ResumeContent, out_path: Path, compression_level: int = 0) -> Path:
    """`compression_level` mirrors app.resume.pdf_writer.write_pdf's ladder --
    reduces base body font size only (bounded by ONE_PAGE_MIN_FONT_SIZE), so
    the DOCX stays visually consistent with a compressed PDF. Structure is
    never altered here -- content-level compression (removing/shortening
    bullets, skills, summary) happens once, upstream, on the shared
    ResumeContent before either writer is called."""
    level = _clamp(compression_level)

    from app import config

    min_font = config.ONE_PAGE_MIN_FONT_SIZE
    body_size = max(min_font, _BASE_FONT_SIZE - 0.3 * level)

    doc = Document()
    doc.styles["Normal"].font.size = Pt(body_size)

    title = doc.add_heading(resume.full_name or "NEEDS_USER_INPUT", level=0)
    title.alignment = 1

    contact_line = " | ".join(
        filter(None, [resume.email, resume.phone, resume.location, resume.linkedin_url, resume.github_url, resume.portfolio_url])
    )
    p = doc.add_paragraph(contact_line)
    p.alignment = 1

    if resume.target_role:
        target = doc.add_paragraph()
        target.alignment = 1
        run = target.add_run(resume.target_role)
        run.bold = True
        run.font.size = Pt(max(min_font, body_size + 1))

    doc.add_heading("Summary", level=1)
    doc.add_paragraph(resume.summary)

    doc.add_heading("Skills", level=1)
    doc.add_paragraph(", ".join(resume.skills_ordered) if resume.skills_ordered else "NEEDS_USER_INPUT")

    if resume.experience:
        doc.add_heading("Experience", level=1)
        for e in resume.experience:
            heading = doc.add_paragraph()
            run = heading.add_run(f"{e.title} — {e.company}")
            run.bold = True
            run.font.size = Pt(max(min_font, body_size + 0.5))
            doc.add_paragraph(f"{e.start_date} - {e.end_date} | {e.location}")
            for b in e.bullets:
                doc.add_paragraph(b, style="List Bullet")

    if resume.projects:
        doc.add_heading("Projects", level=1)
        for proj in resume.projects:
            heading = doc.add_paragraph()
            run = heading.add_run(proj.name)
            run.bold = True
            if proj.url:
                doc.add_paragraph(proj.url)
            doc.add_paragraph(proj.description)
            for b in proj.bullets:
                doc.add_paragraph(b, style="List Bullet")

    if resume.education:
        doc.add_heading("Education", level=1)
        for ed in resume.education:
            doc.add_paragraph(f"{ed.degree} in {ed.field_of_study} — {ed.school} ({ed.graduation_date})")

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path
