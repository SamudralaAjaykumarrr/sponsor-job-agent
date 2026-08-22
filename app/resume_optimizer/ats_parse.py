"""ATS parseability validation (CLAUDE.md Phase 14 sections 28-32, 62).
Never depends on a proprietary ATS's real scoring engine -- only checks that
our own generated DOCX/PDF/TXT artifacts extract cleanly with the expected
text present, in a sensible order, via widely-used parsing libraries
(python-docx / pypdf)."""

from dataclasses import dataclass, field
from pathlib import Path

from docx import Document
from pypdf import PdfReader

from app.resume.generator import ResumeContent
from app.resume_optimizer.models import ATSParseStatus

PARSER_VERSION = "ats-parse-v1"


@dataclass
class ParseCheckResult:
    status: ATSParseStatus
    reasons: list[str] = field(default_factory=list)
    extracted_text_preview: str = ""

    def as_dict(self) -> dict:
        return {"status": self.status.value, "reasons": self.reasons}


def _expected_terms(resume: ResumeContent) -> list[tuple[str, str]]:
    """(label, term) pairs that must be findable in extracted text for a
    resume to be considered ATS-parseable -- CLAUDE.md section 29."""
    terms = []
    if resume.full_name and resume.full_name != "NEEDS_USER_INPUT":
        terms.append(("candidate name", resume.full_name))
    if resume.email and resume.email != "NEEDS_USER_INPUT":
        terms.append(("contact email", resume.email))
    terms.append(("summary section", "Summary"))
    terms.append(("skills section", "Skills"))
    for e in resume.experience[:1]:
        terms.append(("employer", e.company))
        terms.append(("title", e.title))
    for ed in resume.education[:1]:
        terms.append(("education school", ed.school))
    for p in resume.projects[:1]:
        terms.append(("project", p.name))
    return terms


def _score(extracted_text: str, resume: ResumeContent) -> ParseCheckResult:
    text_lower = extracted_text.lower()
    reasons: list[str] = []
    missing = 0
    for label, term in _expected_terms(resume):
        if not term:
            continue
        if term.lower() not in text_lower:
            reasons.append(f"missing expected text: {label} ('{term}')")
            missing += 1

    if not extracted_text.strip():
        return ParseCheckResult(status=ATSParseStatus.FAIL, reasons=["no text could be extracted from the document"])

    # Reading-order sanity check: the candidate's name should appear before
    # the skills section in extracted order (CLAUDE.md section 31).
    name_idx = text_lower.find((resume.full_name or "").lower()) if resume.full_name and resume.full_name != "NEEDS_USER_INPUT" else -1
    skills_idx = text_lower.find("skills")
    if name_idx != -1 and skills_idx != -1 and name_idx > skills_idx:
        reasons.append("reading order unexpected: skills section appears before candidate name")

    if missing >= 3:
        status = ATSParseStatus.FAIL
    elif missing > 0 or any("reading order" in r for r in reasons):
        status = ATSParseStatus.WARN
    else:
        status = ATSParseStatus.PASS
    return ParseCheckResult(status=status, reasons=reasons, extracted_text_preview=extracted_text[:400])


def validate_docx(path: Path, resume: ResumeContent) -> ParseCheckResult:
    try:
        doc = Document(str(path))
    except Exception as exc:  # noqa: BLE001 -- a malformed docx is a FAIL, not a crash
        return ParseCheckResult(status=ATSParseStatus.FAIL, reasons=[f"could not open DOCX: {exc}"])
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return _score("\n".join(parts), resume)


def validate_pdf(path: Path, resume: ResumeContent) -> ParseCheckResult:
    try:
        reader = PdfReader(str(path))
        text = "\n".join((page.extract_text() or "") for page in reader.pages)
    except Exception as exc:  # noqa: BLE001
        return ParseCheckResult(status=ATSParseStatus.FAIL, reasons=[f"could not open/extract PDF: {exc}"])
    return _score(text, resume)


def validate_txt(path: Path, resume: ResumeContent) -> ParseCheckResult:
    try:
        text = path.read_text()
    except Exception as exc:  # noqa: BLE001
        return ParseCheckResult(status=ATSParseStatus.FAIL, reasons=[f"could not read TXT: {exc}"])
    return _score(text, resume)


@dataclass
class ATSParseReport:
    docx: ParseCheckResult
    pdf: ParseCheckResult
    txt: ParseCheckResult

    @property
    def overall(self) -> ATSParseStatus:
        statuses = [self.docx.status, self.pdf.status, self.txt.status]
        if ATSParseStatus.FAIL in statuses:
            return ATSParseStatus.FAIL
        if ATSParseStatus.WARN in statuses:
            return ATSParseStatus.WARN
        return ATSParseStatus.PASS

    def as_dict(self) -> dict:
        return {
            "overall": self.overall.value,
            "docx": self.docx.as_dict(),
            "pdf": self.pdf.as_dict(),
            "txt": self.txt.as_dict(),
            "parser_version": PARSER_VERSION,
        }


def validate_all(docx_path: Path, pdf_path: Path, txt_path: Path, resume: ResumeContent) -> ATSParseReport:
    return ATSParseReport(
        docx=validate_docx(docx_path, resume),
        pdf=validate_pdf(pdf_path, resume),
        txt=validate_txt(txt_path, resume),
    )
