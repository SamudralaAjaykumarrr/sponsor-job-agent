"""Resume layout/visual regression checks (CLAUDE.md Phase 15 sections 34-36).

Deliberately separate from app.resume_optimizer.ats_parse: that module
answers "can a parser extract the expected TEXT from this artifact"
(CLAUDE.md Phase 14 sections 28-32); this module answers a different
question -- "does the generated LAYOUT look structurally sound" -- page
count, blank pages, visible headings, bullets actually rendered as bullets,
contact info near the top. Both are real checks a resume artifact should
pass; neither replaces the other, and callers that want the full picture
run both (see tests/test_resume_visual_regression.py).

Does not rasterize/screenshot pages (no extra heavy rendering dependency
added for this) -- structural inspection via python-docx (paragraph styles)
and pypdf (per-page extracted text length, page count, page geometry) is
"practical" per CLAUDE.md Phase 15 section 34 without requiring a new
dependency. Never enforces an arbitrary one-page rule (section 36): a long
page count is a WARN, never a FAIL, since legitimate content can need more
than one page."""

from dataclasses import dataclass, field
from pathlib import Path

from docx import Document
from pypdf import PdfReader

CHECK_VERSION = "visual-regression-v1"

# A resume this long is almost certainly a generation bug (runaway
# duplication, an infinite bullet loop) rather than legitimately dense
# content -- flagged as a warning, never a hard failure.
PATHOLOGICAL_PAGE_COUNT = 4

# Below this many extracted characters, a PDF page is treated as
# effectively blank (allows for a stray page number/footer without being a
# real content page).
_BLANK_PAGE_CHAR_THRESHOLD = 20


@dataclass
class VisualCheckResult:
    status: str  # "PASS" | "WARN" | "FAIL"
    reasons: list[str] = field(default_factory=list)
    page_count: int = 0

    def as_dict(self) -> dict:
        return {"status": self.status, "reasons": self.reasons, "page_count": self.page_count}


def check_pdf_layout(path: Path) -> VisualCheckResult:
    """No clipped text / reasonable page count / no blank pages -- the PDF-
    specific subset of CLAUDE.md section 34's checklist (headings/bullets
    aren't meaningfully distinguishable from pypdf's flat text extraction,
    so those are checked on the DOCX artifact instead -- both are always
    generated from the same ResumeContent, so a DOCX-level structural
    problem reflects the same underlying content problem)."""
    try:
        reader = PdfReader(str(path))
    except Exception as exc:  # noqa: BLE001
        return VisualCheckResult("FAIL", [f"could not open PDF: {exc}"])

    page_count = len(reader.pages)
    if page_count == 0:
        return VisualCheckResult("FAIL", ["PDF has zero pages"], 0)

    reasons: list[str] = []
    blank_pages = []
    for i, page in enumerate(reader.pages):
        text = (page.extract_text() or "").strip()
        if len(text) < _BLANK_PAGE_CHAR_THRESHOLD:
            blank_pages.append(i + 1)

    status = "PASS"
    if blank_pages:
        reasons.append(f"page(s) with little/no extractable text (likely blank): {blank_pages}")
        status = "FAIL" if len(blank_pages) == page_count else "WARN"

    if page_count >= PATHOLOGICAL_PAGE_COUNT:
        reasons.append(f"page count is {page_count} -- unusually long for a resume; verify content isn't "
                        "overflowing uncontrolled (this is a warning, not a hard one-page requirement)")
        if status == "PASS":
            status = "WARN"

    return VisualCheckResult(status, reasons, page_count)


def check_docx_structure(path: Path, resume) -> VisualCheckResult:
    """Headings visible / bullets render / contact info readable / no
    overlapping content (python-docx has no native overlap concept -- a
    document built purely from add_heading()/add_paragraph() calls, as this
    project's writer does, structurally cannot overlap two blocks; this
    check instead verifies every expected block actually produced its own
    paragraph, which is the failure mode a real content-generation bug
    would show up as)."""
    try:
        doc = Document(str(path))
    except Exception as exc:  # noqa: BLE001
        return VisualCheckResult("FAIL", [f"could not open DOCX: {exc}"])

    reasons: list[str] = []
    paragraphs = doc.paragraphs
    non_empty = [p for p in paragraphs if p.text.strip()]
    if not non_empty:
        return VisualCheckResult("FAIL", ["DOCX has no visible paragraph text"], 0)

    heading_styles = [p for p in paragraphs if p.style is not None and p.style.name.startswith("Heading")]
    if not heading_styles:
        reasons.append("no paragraphs use a Heading style -- section headings may not be visually distinguishable")

    bullet_expected = bool(
        any(e.bullets for e in getattr(resume, "experience", []))
        or any(p.bullets for p in getattr(resume, "projects", []))
    )
    bullet_styles = [p for p in paragraphs if p.style is not None and p.style.name == "List Bullet"]
    if bullet_expected and not bullet_styles:
        reasons.append("resume content includes bullets but no paragraph uses the 'List Bullet' style")

    top_text = "\n".join(p.text for p in non_empty[:5])
    contact_fields = [f for f in (getattr(resume, "email", ""), getattr(resume, "phone", "")) if f and f != "NEEDS_USER_INPUT"]
    missing_contact = [f for f in contact_fields if f not in top_text]
    if missing_contact:
        reasons.append(f"contact field(s) not found near the top of the document: {missing_contact}")

    status = "FAIL" if missing_contact else ("WARN" if reasons else "PASS")
    return VisualCheckResult(status, reasons, page_count=0)


@dataclass
class VisualRegressionReport:
    pdf: VisualCheckResult
    docx: VisualCheckResult

    @property
    def overall(self) -> str:
        statuses = [self.pdf.status, self.docx.status]
        if "FAIL" in statuses:
            return "FAIL"
        if "WARN" in statuses:
            return "WARN"
        return "PASS"

    def as_dict(self) -> dict:
        return {"overall": self.overall, "pdf": self.pdf.as_dict(), "docx": self.docx.as_dict(), "check_version": CHECK_VERSION}


def validate_layout(pdf_path: Path, docx_path: Path, resume) -> VisualRegressionReport:
    return VisualRegressionReport(pdf=check_pdf_layout(pdf_path), docx=check_docx_structure(docx_path, resume))
